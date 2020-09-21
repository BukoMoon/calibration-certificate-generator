#! python3

from os import path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt

# Example variables
img = 'HLP-Logo-Aus.png'
set_date = '3/9/20'
sn = 'HLP–PTPB005457'
model = 'Pocket Temp Pro'
name = 'L.Adams'
months = ' 12 '

# Text to be displayed in a conformance certificate
heading_text = '\rMODEL: ' + model + '\r\rCONFORMANCE CERTIFICATE'
paragraph_text = '\rSerial Number:  ' + sn + '\r\rThis unit has had an operational and calibration' + \
                 '\r\rcheck on  ' + set_date + '\r\r& meets the specifications as set out in the Manufacturers ' + \
                 'specification sheet included with the unit & meets the Australian Food Standards requirements' + \
                 '\r\rThis certificate is valid for' + months + 'months from the above date.\r'

# Text to be displayed for the single conformance certificate
compliance_heading = '\rCertificate of Compliance\rModel:  ' + model
# Add valid_text in-between compliance_heading and compliance_text
compliance_text = '\rThe above unit meets the stated specifications as set out in the Manufacturers specification ' + \
                  'sheet included with the unit & meets the Australian Food Standards requirements.\r\r\r\r\r\r' + \
                  'Signed\t\t' + name + '\t\t\t\tDate Issued: ' + set_date

# Text to be displayed in a calibration certificate
heading_text2 = '\rCalibration Certificate'
valid_text = '(This Certificate valid for ' + months + ' Months)\r\r\r'
paragraph_text2 = '\rChecked against NATA calibrated unit – AZ8801\rS/N  ' + sn + 'Calibration Report: 41598-4\r\r\r' +\
                  'The above unit meets the stated specifications as set out in the Manufacturers specification ' + \
                  'sheet included with the unit & meets the Australian Food Standards requirements.\r\r\r\r' + \
                  'Signed\t' + name + '\t\tDate Issued: ' + set_date


class Template(object):

    def __init__(self, date, serial_number, save_location):
        """

        :param date: The current or date set by the user
        :param serial_number: The serial number of the device
        :param save_location: The name of the file you wish to save the docx as
        """
        # Initialising the path for where the document will be saved
        dir_path = path.dirname(__file__)
        self.date = date
        self.serial_number = serial_number
        self.save_location = path.join(dir_path, save_location)
        self.create_docx()

    def create_docx(self):
        doc = docx.Document()

        doc.save(self.save_location)

    def create_table(self, doc, rows=0, cols=0):
        """

        :param doc: Document to create the table on
        :param rows: How many rows across the table has
        :param cols: How many columns down the table has
        :return: returns the table it creates
        """
        table = doc.add_table(rows, cols)
        table.cell(0, 0).width = Cm(9)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'TableGrid'
        table.autofit = False
        table.width = Cm(9)
        return table

    @staticmethod
    def margin_size(doc, size):
        """

        :param doc: The document you wish to change the margins for
        :param size: int - size of the margin in centimetres
        :return: None
        """
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(size)
            section.bottom_margin = Cm(size)
            section.left_margin = Cm(size)
            section.right_margin = Cm(size)

    @staticmethod
    def paragraph_runs(doc, *args: object, font: object = 'Calibri') -> object:
        """

        :param doc: A document or a table to add each paragraph/run to
        :param args: Any text that you wish to be displayed in the document or table
        :param font: The name of the font you wish to use for each paragraph
        :return: List of each run that was made to edit further if need be. e.g font size
        """
        list_of_runs = []
        for arg in args:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(arg)
            run.font.name = font
            list_of_runs.append(run)

        return list_of_runs

    @staticmethod
    def center_all_paragraphs(doc):
        """

        :param doc: The document or a table you want the paragraphs centered for
        :return: None
        """
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def font_size(*paragraphs, size=12):
        """

        :param paragraphs: The paragraphs you want to change the size of
        :param size: The size of the font in points
        :return:
        """
        for paragraph in paragraphs:
            paragraph.font.size = Pt(size)

    @staticmethod
    def bold(*args):
        """

        :param args: Any paragraph run that you want to make bold.
        :return: None
        """
        for arg in args:
            arg.bold = True

    @staticmethod
    def underline(*args):
        """

        :param args: Any paragraph run that you want to have an underline.
        :return: None
        """
        for arg in args:
            arg.underline = True


class CalibrationCertificate(Template):

    def __init__(self, date, serial_number, save_location):
        super().__init__(date, serial_number, save_location)

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 2)
        self.docx_contents(doc)
        doc.save(self.save_location)

    def docx_contents(self, doc):
        doc.add_picture(img, width=Cm(17.2), height=Cm(3.95))
        heading_run, valid_run, text_run = self.paragraph_runs(doc, heading_text, valid_text, paragraph_text,
                                                               font='Times New Roman')
        self.center_all_paragraphs(doc)

        self.font_size(heading_run, size=28)
        self.font_size(valid_run, size=10)
        self.font_size(text_run, size=22)


class SingleConformance(Template):

    def __init__(self, date, serial_number, save_location):
        super().__init__(date, serial_number, save_location)

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 2)
        self.docx_contents(doc)
        doc.save(self.save_location)

    def docx_contents(self, doc):
        doc.add_picture(img, width=Cm(17.3), height=Cm(3.95))
        heading_run, valid_run, text_run = self.paragraph_runs(doc, heading_text2, valid_text, paragraph_text2,
                                                               font='Times New Roman')
        self.center_all_paragraphs(doc)

        # Font sizes for each paragraph
        self.font_size(heading_run, size=28)
        self.font_size(valid_run, size=10)
        self.font_size(text_run, size=22)


class MultipleConformance(Template):

    def __init__(self, date, serial_number, save_location):
        super().__init__(date, serial_number, save_location)

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 1)
        self.create_table(doc, rows=2, cols=2)

        doc.save(self.save_location)

    def create_table(self, doc, rows=0, cols=0):
        table = super().create_table(doc, rows=rows, cols=cols)
        self.table_contents(table, rows)
        return table

    def table_contents(self, table, rows):
        for i in range(0, rows):
            for cell in range(0, rows):
                table_rows = table.rows[i].cells[cell]
                p = table_rows.add_paragraph()
                p.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                p.height = Cm(11)
                table.rows[i].height = Cm(9)

                # Adding image
                p.add_run().add_picture(img, width=Cm(9.09), height=Cm(2.41))
                # Adding runs to the paragraph
                heading, para = self.paragraph_runs(table_rows, heading_text, paragraph_text, font='Arial')
                self.center_all_paragraphs(table_rows)
                self.bold(heading)
                self.underline(heading)
                self.font_size(heading, size=14)
                self.font_size(para, size=12)


if __name__ == '__main__':
    # Generated Calibration Certificate must have 'device_model and serial number' for the name.
    MultipleConformance(set_date, sn, 'generated.docx')
    CalibrationCertificate(set_date, sn, 'generated-calibration.docx')
    SingleConformance(set_date, sn, 'generated-conformance-single.docx')
