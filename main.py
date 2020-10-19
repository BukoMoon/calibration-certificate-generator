#! python3
import datetime
import sys
import docx
import win32api
import win32print


from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from PyQt5 import QtWidgets, QtCore, QtGui
from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QApplication, QMainWindow
from os import path


class Template(object):

    def __init__(self, file_name):
        """

        :param file_name: The name of the file you wish to save the docx as
        """
        # Initialising the path for where the document will be saved
        dir_path = path.dirname(__file__)
        self.save_location = path.join(dir_path, file_name)
        self._model = ''
        self._date = ''
        self._prefix = ''
        self._name = ''
        self._months = ''
        self.check = ''
        self.temp = ''
        self._serial_number = None

    def create_docx(self):
        doc = docx.Document()

        doc.save(self.save_location)

    def create_table(self, doc, rows=0, cols=0):
        """

        :param doc: Document to create the table on
        :param rows: How many rows across the table has
        :param cols: How many columns down the table has
        :return: returns the table it creates
        """
        table = doc.add_table(rows, cols)
        table.cell(0, 0).width = Cm(9)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'TableGrid'
        table.autofit = False
        table.width = Cm(9)
        return table

    @staticmethod
    def margin_size(doc, size):
        """

        :param doc: The document you wish to change the margins for
        :param size: int - size of the margin in centimetres
        :return: None
        """
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(size)
            section.bottom_margin = Cm(size)
            section.left_margin = Cm(size)
            section.right_margin = Cm(size)

    @staticmethod
    def paragraph_runs(doc, *args: object, font: object = 'Calibri') -> object:
        """

        :param doc: A document or a table to add each paragraph/run to
        :param args: Any text that you wish to be displayed in the document or table
        :param font: The name of the font you wish to use for each paragraph
        :return: List of each run that was made to edit further if need be. e.g font size
        """
        list_of_runs = []
        for arg in args:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(arg)
            run.font.name = font
            list_of_runs.append(run)

        return list_of_runs

    @staticmethod
    def center_all_paragraphs(doc):
        """

        :param doc: The document or a table you want the paragraphs centered for
        :return: None
        """
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def font_size(*paragraphs, size=12):
        """

        :param paragraphs: The paragraphs you want to change the size of
        :param size: The size of the font in points
        :return: None
        """
        for paragraph in paragraphs:
            paragraph.font.size = Pt(size)

    @staticmethod
    def bold(*args):
        """

        :param args: Any paragraph run that you want to make bold.
        :return: None
        """
        for arg in args:
            arg.bold = True

    @staticmethod
    def underline(*args):
        """

        :param args: Any paragraph run that you want to have an underline.
        :return: None
        """
        for arg in args:
            arg.underline = True

    @property
    def model(self) -> str:
        return self._model

    @model.setter
    def model(self, device_model: str):
        self._model = device_model

    @property
    def date(self) -> str:
        return self._date

    @date.setter
    def date(self, set_date: str):
        self._date = set_date

    @property
    def serial_number(self) -> int:
        return self._serial_number

    @serial_number.setter
    def serial_number(self, device_serial_number: str):
        self._serial_number = int(device_serial_number)

    @property
    def months(self) -> str:
        return self._months

    @months.setter
    def months(self, valid_time: str):
        self._months = valid_time

    @property
    def name(self) -> str:
        return self._name

    @name.setter
    def name(self, set_name: str):
        self._name = set_name

    @property
    def prefix(self) -> str:
        return self._prefix

    @prefix.setter
    def prefix(self, value=''):
        if self.model.lower() == 'pocket temp pro':
            self._prefix = 'HLP-PTP'
        elif self.model.lower() == 'pocket temp blue':
            self._prefix = 'HLP-PTPB'
        else:
            self._prefix = value

    @property
    def check(self) -> str:
        return str(self._check)

    @check.setter
    def check(self, temperature_value: float):
        self._check = temperature_value

    @property
    def temp(self) -> str:
        return str(self._temp)

    @temp.setter
    def temp(self, temp_value: float):
        self._temp = temp_value


class CalibrationCertificate(Template):

    def __init__(self, file_name):
        super().__init__(file_name)
        self.img = 'HLP-Logo-Aus.png'

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 2)
        self.docx_contents(doc)
        doc.save(self.save_location)

    def docx_contents(self, doc):
        doc.add_picture(self.img, width=Cm(17.2), height=Cm(3.95))
        # Text to be displayed in a calibration certificate
        heading_text = 'Calibration Certificate'
        valid_text = '(This Certificate valid for ' + self.months + ' Months)\r'
        paragraph_text = 'Model:\t' + self.model + \
                         '\r SN: ' + self.prefix + str(self.serial_number).zfill(6) + '\rRead ' + self.check + \
                         '°C @ ' + self.temp + ' °C\r\rChecked against NATA calibrated unit AZ8801\r S/N - 9000782' \
                                               'Calibration Report: 41598-4\r\r\rThe above unit meets the stated specifications ' \
                                               'as set out in the Manufacturers specification sheet included with the unit & meets ' \
                                               'the Australian Food Standards requirements.\r\r\rSigned\t' + self.name + \
                         '\t\t\tDate Issued: ' + self.date
        heading_run, valid_run, text_run = self.paragraph_runs(doc, heading_text, valid_text,
                                                               paragraph_text, font='Times New Roman')
        self.center_all_paragraphs(doc)

        self.font_size(heading_run, size=28)
        self.font_size(valid_run, size=10)
        self.font_size(text_run, size=22)


class SingleConformance(Template):

    def __init__(self, file_name):
        super().__init__(file_name)
        self.img = 'HLP-Logo-Aus.png'

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 2)
        self.docx_contents(doc)
        doc.save(self.save_location)

    def docx_contents(self, doc):
        doc.add_picture(self.img, width=Cm(17.3), height=Cm(3.95))
        # Text to be displayed for the single conformance certificate
        compliance_heading = '\rCertificate of Compliance\rModel:  ' + self.model
        valid_text = '(This Certificate valid for ' + self.months + ' Months)\r\r'
        compliance_text = '\rThe above unit meets the stated specifications as set out in the Manufacturers ' \
                          'specification sheet included with the unit & meets the Australian Food Standards ' \
                          'requirements.\r\r\r\r\r' + 'Signed\t\t' + self.name + '\t\t\t\tDate Issued: ' + self.date
        heading_run, valid_run, text_run = self.paragraph_runs(doc, compliance_heading, valid_text,
                                                               compliance_text, font='Times New Roman')
        self.center_all_paragraphs(doc)

        # Font sizes for each paragraph
        self.font_size(valid_run, size=10)
        self.font_size(heading_run, size=28)
        self.font_size(text_run, size=22)


class MultipleConformance(Template):

    def __init__(self, file_name):
        super().__init__(file_name)

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 1)
        self.create_table(doc, rows=2, cols=2)

        doc.save(self.save_location)

    def create_table(self, doc, rows=0, cols=0):
        table = super().create_table(doc, rows=rows, cols=cols)
        self.table_contents(table, rows)
        return table

    def table_contents(self, table, rows):
        img = 'HLP-Logo-Aus.png'

        # Text to be displayed in a conformance certificate
        heading_text = '\rMODEL: ' + self.model + '\r\rCONFORMANCE CERTIFICATE'

        for i in range(0, rows):
            for cell in range(0, rows):
                paragraph_text = '\rSerial Number:  ' + self.prefix + str(self.serial_number).zfill(6) + '\r\rThis ' \
                                                                                                         'unit has had an operational and calibration\rcheck on  ' + self.date + \
                                 '\r\r& meets the specifications as set out in the Manufacturers specification sheet ' \
                                 'included with the unit & meets the Australian Food Standards requirements\r\rThis ' \
                                 'certificate is valid for ' + self.months + ' months from the above date.\r' + \
                                 '\rSigned\t\t' + self.name + '\r'
                table_rows = table.rows[i].cells[cell]
                p = table_rows.add_paragraph()
                p.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                p.height = Cm(11)
                table.rows[i].height = Cm(9)

                # Adding image
                p.add_run().add_picture(img, width=Cm(9.09), height=Cm(2.41))
                # Adding runs to the paragraph
                heading, para = self.paragraph_runs(table_rows, str(heading_text), str(paragraph_text), font='Arial')
                self.center_all_paragraphs(table_rows)
                self.bold(heading)
                self.underline(heading)
                self.font_size(heading, size=14)
                self.font_size(para, size=12)
                self.serial_number = self.serial_number + 1


class GeneratorPrinter(object):

    def __init__(self, filename):
        self.filename = filename

    def print_docx(self):
        win32api.ShellExecute(0, "print", self.filename, '/d:"%s"' % win32print.GetDefaultPrinter(), ".", 0)


class MyWindow(QMainWindow):

    def __init__(self):
        super(MyWindow, self).__init__()
        self.setGeometry(200, 200, 800, 600)
        self.setWindowTitle('HLP Controls Pty Ltd - Certificate Generator - By Louis Adams')
        _translate = QtCore.QCoreApplication.translate

        self._print = None

        # Initialising Widgets
        self.central_widget = QtWidgets.QWidget(self)

        # Creating groups
        self.certificate = QtWidgets.QButtonGroup()
        self.months = QtWidgets.QButtonGroup()

        # Toolbar
        self.menu_bar = QtWidgets.QMenuBar(self)
        self.menu_file = QtWidgets.QMenu(self.menu_bar)
        self.status_bar = QtWidgets.QStatusBar(self)
        self.action_exit = QtWidgets.QAction(self)
        self.init_toolbar(_translate)

        # Certificate checkboxes
        self.check_box = QtWidgets.QCheckBox(self.central_widget)
        self.check_box_2 = QtWidgets.QCheckBox(self.central_widget)
        self.check_box_5 = QtWidgets.QCheckBox(self.central_widget)
        self.certificate_checkboxes(_translate)

        self.label_checked_at = QtWidgets.QLabel(self.central_widget)
        self.label_temp = QtWidgets.QLabel(self.central_widget)
        self.temperature = QtWidgets.QLineEdit(self.central_widget)
        self.checked = QtWidgets.QLineEdit(self.central_widget)
        self.calibration_lines_and_labels(_translate)

        # Months checkboxes
        self.check_box_3 = QtWidgets.QCheckBox(self.central_widget)
        self.check_box_4 = QtWidgets.QCheckBox(self.central_widget)
        self.month_checkboxes(_translate)

        # Line edit boxes - Model, Serial Number, Name & Date
        self.text_edit = QtWidgets.QLineEdit(self.central_widget)
        self.text_edit_2 = QtWidgets.QLineEdit(self.central_widget)
        self.text_edit_3 = QtWidgets.QLineEdit(self.central_widget)
        self.date_edit = QtWidgets.QDateEdit(self.central_widget)
        self.text_date_widgets(_translate)

        # Generate buttons
        self.button1 = QtWidgets.QPushButton(self.central_widget)
        self.button1_2 = QtWidgets.QPushButton(self.central_widget)
        self.generate_buttons(_translate)

        # Labels
        self.label = QtWidgets.QLabel(self.central_widget)
        self.label1 = QtWidgets.QLabel(self.central_widget)
        self.label_2 = QtWidgets.QLabel(self.central_widget)
        self.label_3 = QtWidgets.QLabel(self.central_widget)
        self.label_4 = QtWidgets.QLabel(self.central_widget)
        self.label_5 = QtWidgets.QLabel(self.central_widget)
        self.label_6 = QtWidgets.QLabel(self.central_widget)
        self.labels(_translate)

        self.init_ui()

    def init_toolbar(self, _translate):
        self.menu_file.setTitle(_translate("MainWindow", "File"))
        self.action_exit.setText(_translate("MainWindow", "Exit"))
        self.action_exit.setStatusTip(_translate("MainWindow", "Exit the program"))
        self.action_exit.setShortcut(_translate("MainWindow", "Ctrl+E"))

        self.menu_bar.setObjectName("menu_bar")
        self.menu_file.setObjectName("menu_file")
        self.setMenuBar(self.menu_bar)
        self.status_bar.setObjectName("status_bar")
        self.setStatusBar(self.status_bar)
        self.menu_file.addAction(self.action_exit)
        self.menu_bar.addAction(self.menu_file.menuAction())

    def calibration_lines_and_labels(self, _translate):
        self.label_temp.setText(_translate("MainWindow", "Temperature Read"))
        self.label_checked_at.setText(_translate("MainWindow", "@"))

        self.temperature.setPlaceholderText(_translate("MainWindow", "Enter temperature here on thermometer"))
        self.temperature.setObjectName("temperature")
        self.temperature.setMaxLength(5)

        self.checked.setPlaceholderText(_translate("MainWindow", "Enter temperature from master thermometer here"))
        self.temperature.setObjectName("checked against temperature")
        self.temperature.setMaxLength(5)

    def certificate_checkboxes(self, _translate):
        # Adding checkboxes to a group
        self.certificate.addButton(self.check_box_2)
        self.certificate.addButton(self.check_box)
        self.certificate.addButton(self.check_box_5)
        self.check_box.isChecked = True

        # Setting name & text
        self.check_box.setObjectName("check_box")
        self.check_box_2.setObjectName("check_box_2")
        self.check_box_5.setObjectName("check_box_5")
        self.check_box.setText(_translate("MainWindow", "Single Conformance Certificate"))
        self.check_box_2.setText(_translate("MainWindow", "Calibration Certificate"))
        self.check_box_5.setText(_translate("MainWindow", "Multiple Conformance Certificates"))

    def month_checkboxes(self, _translate):
        # Adding checkboxes to a group
        self.months.addButton(self.check_box_3)
        self.months.addButton(self.check_box_4)
        self.check_box_3.isChecked = True

        # Setting name & text
        self.check_box_3.setText(_translate("MainWindow", "12 Months"))
        self.check_box_4.setText(_translate("MainWindow", "24 Months"))
        self.check_box_3.setObjectName("check_box_3")
        self.check_box_4.setObjectName("check_box_4")

    def text_date_widgets(self, _translate):
        self.text_edit.setToolTip(_translate("MainWindow", "Enter your name here"))
        self.text_edit.setPlaceholderText(_translate("MainWindow", "Enter your name here E.g L.Adams"))
        self.text_edit_2.setToolTip(_translate("MainWindow", "Enter the Device\'s Serial Number"))
        self.text_edit_2.setPlaceholderText(_translate("MainWindow", "Enter the Device Serial Number Here"))
        self.text_edit_3.setToolTip(_translate("MainWindow", "Enter your Device Model Here"))
        self.text_edit_3.setPlaceholderText(_translate("MainWindow", "Enter the Device Model Here"))
        self.date_edit.setToolTip(_translate("MainWindow", "The date on the certificate - normally the current date"))
        self.date_edit.setDisplayFormat(_translate("MainWindow", "dd/MM/yyyy"))

        self.text_edit.setObjectName("text_edit")
        self.text_edit.setMaxLength(20)
        self.text_edit_2.setObjectName("text_edit_2")
        self.text_edit_2.setMaxLength(20)
        self.text_edit_3.setObjectName("text_edit_3")
        self.text_edit_3.setMaxLength(20)
        self.date_edit.setObjectName("date_edit")

    def generate_buttons(self, _translate):
        self.button1.setToolTip(_translate("MainWindow", "Click to generate your certificate/s"))
        self.button1.setText(_translate("MainWindow", "Generate and Print"))
        self.button1_2.setToolTip(_translate("MainWindow", "Click to generate your certificate/s"))
        self.button1_2.setText(_translate("MainWindow", "Generate"))

        self.button1.setObjectName("button1")
        self.button1_2.setObjectName("button1_2")

    def labels(self, _translate):
        self.label.setText(_translate("MainWindow", "Date:"))
        self.label1.setText(_translate("MainWindow", "Select a Template"))
        self.label_2.setText(_translate("MainWindow", "Enter your name below"))
        self.label_3.setText(_translate("MainWindow", "Device Serial Number"))
        self.label_4.setText(_translate("MainWindow", "Certificate Validation Period"))
        self.label_5.setText(_translate("MainWindow", "Device Model"))

    def init_ui(self):
        self.resize(900, 700)
        self.setMinimumSize(QtCore.QSize(900, 700))
        self.setCentralWidget(self.central_widget)

        # Initialising font sizes
        small_font = QtGui.QFont()
        small_font.setPointSize(11)
        font = QtGui.QFont()
        font.setPointSize(14)

        self.central_widget.setObjectName("central_widget")
        self.menu_bar.setGeometry(QtCore.QRect(0, 0, 900, 22))

        # Certificate checkboxes locations
        self.check_box.setGeometry(QtCore.QRect(80, 250, 241, 31))
        self.check_box_2.setGeometry(QtCore.QRect(80, 350, 171, 23))
        self.check_box_5.setGeometry(QtCore.QRect(80, 300, 261, 31))

        # Month checkboxes locations
        self.check_box_3.setGeometry(QtCore.QRect(80, 440, 121, 31))
        self.check_box_4.setGeometry(QtCore.QRect(80, 480, 111, 31))

        # LineEdit Widgets & Date locations
        self.date_edit.setGeometry(QtCore.QRect(124, 534, 111, 31))
        self.date_edit.setCalendarPopup(True)
        self.date_edit.setDate(QtCore.QDate(2020, 9, 12))
        self.text_edit.setGeometry(QtCore.QRect(500, 420, 321, 31))
        self.text_edit_2.setGeometry(QtCore.QRect(500, 319, 321, 31))
        self.text_edit_3.setGeometry(QtCore.QRect(500, 230, 321, 31))

        self.temperature.setGeometry(QtCore.QRect(280, 500, 120, 20))
        self.checked.setGeometry(QtCore.QRect(280, 540, 120, 20))
        self.label_temp.setFont(small_font)
        self.label_temp.setObjectName("label_temp")
        self.label_temp.setAlignment(QtCore.Qt.AlignCenter)
        self.label_temp.setGeometry(QtCore.QRect(260, 480, 151, 16))
        self.label_checked_at.setFont(small_font)
        self.label_checked_at.setObjectName("label_checked_at")
        self.label_checked_at.setAlignment(QtCore.Qt.AlignCenter)
        self.label_checked_at.setGeometry(QtCore.QRect(320, 520, 31, 16))

        # Button locations
        self.button1.setGeometry(QtCore.QRect(650, 500, 171, 61))
        self.button1_2.setGeometry(QtCore.QRect(490, 500, 151, 61))

        # Label locations & font
        self.label.setGeometry(QtCore.QRect(80, 540, 71, 21))
        self.label.setFont(small_font)
        self.label.setObjectName("label")
        self.label1.setGeometry(QtCore.QRect(80, 210, 171, 31))
        self.label1.setFont(font)
        self.label1.setObjectName("label1")
        self.label_2.setGeometry(QtCore.QRect(500, 390, 321, 31))
        self.label_2.setFont(small_font)
        self.label_2.setAlignment(QtCore.Qt.AlignCenter)
        self.label_2.setObjectName("label_2")
        self.label_3.setGeometry(QtCore.QRect(500, 290, 321, 20))
        self.label_3.setAlignment(QtCore.Qt.AlignCenter)
        self.label_3.setObjectName("label_3")
        self.label_4.setGeometry(QtCore.QRect(80, 390, 251, 41))
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.label_5.setGeometry(QtCore.QRect(500, 200, 321, 31))
        self.label_5.setAlignment(QtCore.Qt.AlignCenter)
        self.label_5.setObjectName("label_5")

        # Logo Image
        self.label_6.setGeometry(QtCore.QRect(40, 10, 801, 151))
        self.label_6.setText("")
        self.label_6.setPixmap(QtGui.QPixmap("HLP-Logo-Aus.jpg"))
        self.label_6.setScaledContents(True)
        self.label_6.setObjectName("label_6")

        self.button1.clicked.connect(self.generate)
        if self.button1_2.clicked:
            self.print_files = True
        self.button1_2.clicked.connect(self.generate)
        QtCore.QMetaObject.connectSlotsByName(self)
        self.check_box_2.stateChanged.connect(self.state_change)
        self.show()

    def state_change(self):
        if self.check_box_2.isChecked():
            self.temperature.show()
            self.checked.show()
            self.label_temp.show()
            self.label_checked_at.show()
        else:
            self.temperature.hide()
            self.checked.hide()
            self.label_temp.hide()
            self.label_checked_at.hide()

    @pyqtSlot()
    def generate(self):
        convert_date = self.date_edit.date().toPyDate()
        set_date = datetime.datetime.strftime(convert_date, '%d/%m/%Y   ')
        set_name = self.text_edit.text()
        set_serial_number = str(self.text_edit_2.text())
        set_model = self.text_edit_3.text()
        if self.label_temp.text or self.label_checked_at.text == "":
            set_checked_at = ""
            set_temperature = ""
        else:
            set_checked_at = self.label_temp.text()
            set_temperature = self.label_checked_at.text()
        file_name = 'generated.docx'

        if self.check_box_3.isChecked:
            set_months = ' 12 '
        else:
            set_months = ' 24 '

        if self.check_box.isChecked:
            single_conformance = SingleConformance(file_name)
            single_conformance.model = set_model
            single_conformance.date = set_date
            single_conformance.name = set_name
            single_conformance.months = set_months
            single_conformance.serial_number = set_serial_number
            single_conformance.prefix = ''
            single_conformance.create_docx()
            if self.print_files:
                self.printer(file_name, False)

        elif self.check_box_2.isChecked:
            calibration_certificate = CalibrationCertificate(file_name)
            calibration_certificate.model = set_model
            calibration_certificate.date = set_date
            calibration_certificate.name = set_name
            calibration_certificate.months = set_months
            calibration_certificate.serial_number = set_serial_number
            calibration_certificate.prefix = ''
            calibration_certificate.check = set_checked_at
            calibration_certificate.temp = set_temperature
            calibration_certificate.create_docx()
            if self.print_files:
                self.printer(file_name, False)

        else:
            self.generate_multiple(set_date, set_serial_number, set_model, set_months, set_name)

    def generate_multiple(self, set_date, set_serial_number, set_model, set_months, set_name):
        print_multiple = True
        for i in range(0, 7):
            print(f'Generating docx {i}')
            multiple_conformance = MultipleConformance('generated' + str(i) + '.docx')
            multiple_conformance.model = set_model
            multiple_conformance.date = set_date
            multiple_conformance.name = set_name
            multiple_conformance.months = set_months
            multiple_conformance.serial_number = set_serial_number
            multiple_conformance.prefix = ''
            multiple_conformance.create_docx()
            set_serial_number += 4
            if print_multiple and self.print_files:
                self.printer('generated' + str(i) + '.docx', True)

    @staticmethod
    def printer(file, multiple_files: bool):
        if multiple_files:
            for file_number in range(0, 7):
                print(f'{file_number + 1}. Printing docx, {file}')
        else:
            print(f'Printing docx, {file}')

    @property
    def print_files(self):
        return self._print

    @print_files.setter
    def print_files(self, print_docx=False):
        self._print = print_docx


def window():
    app = QApplication(sys.argv)
    win = MyWindow()

    win.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    window()
