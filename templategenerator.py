#! python3

from os import path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt


class Template(object):

    def __init__(self, file_name):
        """

        :param file_name: The name of the file you wish to save the docx as
        """
        # Initialising the path for where the document will be saved
        dir_path = path.dirname(__file__)
        self.save_location = path.join(dir_path, file_name)
        self._model = ''
        self._date = ''
        self._prefix = ''
        self._name = ''
        self._months = ''
        self.check = ''
        self.temp = ''
        self._serial_number = None

    def create_docx(self):
        doc = docx.Document()

        doc.save(self.save_location)

    def create_table(self, doc, rows=0, cols=0):
        """

        :param doc: Document to create the table on
        :param rows: How many rows across the table has
        :param cols: How many columns down the table has
        :return: returns the table it creates
        """
        table = doc.add_table(rows, cols)
        table.cell(0, 0).width = Cm(9)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.style = 'TableGrid'
        table.autofit = False
        table.width = Cm(9)
        return table

    @staticmethod
    def margin_size(doc, size):
        """

        :param doc: The document you wish to change the margins for
        :param size: int - size of the margin in centimetres
        :return: None
        """
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(size)
            section.bottom_margin = Cm(size)
            section.left_margin = Cm(size)
            section.right_margin = Cm(size)

    @staticmethod
    def paragraph_runs(doc, *args: object, font: object = 'Calibri') -> object:
        """

        :param doc: A document or a table to add each paragraph/run to
        :param args: Any text that you wish to be displayed in the document or table
        :param font: The name of the font you wish to use for each paragraph
        :return: List of each run that was made to edit further if need be. e.g font size
        """
        list_of_runs = []
        for arg in args:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(arg)
            run.font.name = font
            list_of_runs.append(run)

        return list_of_runs

    @staticmethod
    def center_all_paragraphs(doc):
        """

        :param doc: The document or a table you want the paragraphs centered for
        :return: None
        """
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    @staticmethod
    def font_size(*paragraphs, size=12):
        """

        :param paragraphs: The paragraphs you want to change the size of
        :param size: The size of the font in points
        :return: None
        """
        for paragraph in paragraphs:
            paragraph.font.size = Pt(size)

    @staticmethod
    def bold(*args):
        """

        :param args: Any paragraph run that you want to make bold.
        :return: None
        """
        for arg in args:
            arg.bold = True

    @staticmethod
    def underline(*args):
        """

        :param args: Any paragraph run that you want to have an underline.
        :return: None
        """
        for arg in args:
            arg.underline = True

    @property
    def model(self) -> str:
        return self._model

    @model.setter
    def model(self, device_model: str):
        self._model = device_model

    @property
    def date(self) -> str:
        return self._date

    @date.setter
    def date(self, set_date: str):
        self._date = set_date

    @property
    def serial_number(self) -> int:
        return self._serial_number

    @serial_number.setter
    def serial_number(self, device_serial_number: str):
        self._serial_number = int(device_serial_number)

    @property
    def months(self) -> str:
        return self._months

    @months.setter
    def months(self, valid_time: str):
        self._months = valid_time

    @property
    def name(self) -> str:
        return self._name

    @name.setter
    def name(self, set_name: str):
        self._name = set_name

    @property
    def prefix(self) -> str:
        return self._prefix

    @prefix.setter
    def prefix(self, value=''):
        if self.model.lower() == 'pocket temp pro':
            self._prefix = 'HLP-PTP'
        elif self.model.lower() == 'pocket temp blue':
            self._prefix = 'HLP-PTPB'
        else:
            self._prefix = value

    @property
    def check(self) -> str:
        return str(self._check)

    @check.setter
    def check(self, temperature_value: float):
        self._check = temperature_value

    @property
    def temp(self) -> str:
        return str(self._temp)

    @temp.setter
    def temp(self, temp_value: float):
        self._temp = temp_value


class CalibrationCertificate(Template):

    def __init__(self, file_name):
        super().__init__(file_name)
        self.img = 'HLP-Logo-Aus.png'

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 2)
        self.docx_contents(doc)
        doc.save(self.save_location)

    def docx_contents(self, doc):
        doc.add_picture(self.img, width=Cm(17.2), height=Cm(3.95))
        # Text to be displayed in a calibration certificate
        heading_text = 'Calibration Certificate'
        valid_text = '(This Certificate valid for ' + self.months + ' Months)\r'
        paragraph_text = 'Model:\t' + self.model + \
                         '\r SN: ' + self.prefix + str(self.serial_number).zfill(6) + '\rRead ' + self.check + \
                         '°C @ ' + self.temp + ' °C\r\rChecked against NATA calibrated unit AZ8801\r S/N - 9000782' \
                         'Calibration Report: 41598-4\r\r\rThe above unit meets the stated specifications ' \
                         'as set out in the Manufacturers specification sheet included with the unit & meets ' \
                         'the Australian Food Standards requirements.\r\r\rSigned\t' + self.name + \
                         '\t\t\tDate Issued: ' + self.date
        heading_run, valid_run, text_run = self.paragraph_runs(doc, heading_text, valid_text,
                                                               paragraph_text, font='Times New Roman')
        self.center_all_paragraphs(doc)

        self.font_size(heading_run, size=28)
        self.font_size(valid_run, size=10)
        self.font_size(text_run, size=22)


class SingleConformance(Template):

    def __init__(self, file_name):
        super().__init__(file_name)
        self.img = 'HLP-Logo-Aus.png'

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 2)
        self.docx_contents(doc)
        doc.save(self.save_location)

    def docx_contents(self, doc):
        doc.add_picture(self.img, width=Cm(17.3), height=Cm(3.95))
        # Text to be displayed for the single conformance certificate
        compliance_heading = '\rCertificate of Compliance\rModel:  ' + self.model
        valid_text = '(This Certificate valid for ' + self.months + ' Months)\r\r'
        compliance_text = '\rThe above unit meets the stated specifications as set out in the Manufacturers ' \
                          'specification sheet included with the unit & meets the Australian Food Standards ' \
                          'requirements.\r\r\r\r\r' + 'Signed\t\t' + self.name + '\t\t\t\tDate Issued: ' + self.date
        heading_run, valid_run, text_run = self.paragraph_runs(doc, compliance_heading, valid_text,
                                                               compliance_text, font='Times New Roman')
        self.center_all_paragraphs(doc)

        # Font sizes for each paragraph
        self.font_size(valid_run, size=10)
        self.font_size(heading_run, size=28)
        self.font_size(text_run, size=22)


class MultipleConformance(Template):

    def __init__(self, file_name):
        super().__init__(file_name)

    def create_docx(self):
        doc = docx.Document()
        super().margin_size(doc, 1)
        self.create_table(doc, rows=2, cols=2)

        doc.save(self.save_location)

    def create_table(self, doc, rows=0, cols=0):
        table = super().create_table(doc, rows=rows, cols=cols)
        self.table_contents(table, rows)
        return table

    def table_contents(self, table, rows):
        img = 'HLP-Logo-Aus.png'

        # Text to be displayed in a conformance certificate
        heading_text = '\rMODEL: ' + self.model + '\r\rCONFORMANCE CERTIFICATE'

        for i in range(0, rows):
            for cell in range(0, rows):
                paragraph_text = '\rSerial Number:  ' + self.prefix + str(self.serial_number).zfill(6) + '\r\rThis ' \
                                 'unit has had an operational and calibration\rcheck on  ' + self.date + \
                                 '\r\r& meets the specifications as set out in the Manufacturers specification sheet ' \
                                 'included with the unit & meets the Australian Food Standards requirements\r\rThis ' \
                                 'certificate is valid for ' + self.months + ' months from the above date.\r' + \
                                 '\rSigned\t\t' + self.name + '\r'
                table_rows = table.rows[i].cells[cell]
                p = table_rows.add_paragraph()
                p.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                p.height = Cm(11)
                table.rows[i].height = Cm(9)

                # Adding image
                p.add_run().add_picture(img, width=Cm(9.09), height=Cm(2.41))
                # Adding runs to the paragraph
                heading, para = self.paragraph_runs(table_rows, str(heading_text), str(paragraph_text), font='Arial')
                self.center_all_paragraphs(table_rows)
                self.bold(heading)
                self.underline(heading)
                self.font_size(heading, size=14)
                self.font_size(para, size=12)
                self.serial_number = self.serial_number + 1


if __name__ == '__main__':
    # Date - Serial Number - Model - Months - Name - File Name
    multiple = MultipleConformance('generated.docx')
    multiple.model = 'Pocket Temp Pro'
    multiple.date = '3/9/20'
    multiple.name = 'L.Adams'
    multiple.months = '12'
    multiple.serial_number = 5457
    multiple.prefix = ''
    multiple.create_docx()

    calibration = CalibrationCertificate('generated-calibration.docx')
    calibration.model = 'Pocket Temp Pro'
    calibration.date = '3/9/20'
    calibration.name = 'L.Adams'
    calibration.months = '12'
    calibration.serial_number = 5457
    calibration.prefix = ''
    calibration.check = 24.2
    calibration.temp = 24.1
    calibration.create_docx()

    single = SingleConformance('generated-conformance-single.docx')
    single.model = 'Pocket Temp Pro'
    single.date = '3/9/20'
    single.name = 'L.Adams'
    single.months = '12'
    single.serial_number = 5457
    single.prefix = ''
    single.create_docx()

